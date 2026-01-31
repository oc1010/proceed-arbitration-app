import streamlit as st
from datetime import datetime, date
import vertexai
from vertexai.generative_models import GenerativeModel
from google.oauth2 import service_account
from google.api_core.exceptions import NotFound, Forbidden, ServiceUnavailable, InvalidArgument
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from db import load_complex_data, load_full_config

# ==============================================================================
# 1. HARD MATH ENGINE
# ==============================================================================

def calculate_doc_prod_score(role):
    """Calculates Rejection Rate. Metric: >75% Rejection = 100% Phase Penalty."""
    data = load_complex_data()
    meta = load_full_config().get('meta', {})
    threshold = meta.get('cost_settings', {}).get('doc_prod_threshold', 75.0)
    
    requests = data.get('doc_prod', {}).get(role, [])
    if not requests: return 0.0, False
    
    total = len(requests)
    rejected = sum(1 for r in requests if r.get('status') == 'Denied')
    
    ratio = (rejected / total) * 100 if total > 0 else 0.0
    penalty_triggered = ratio > threshold
    
    return ratio, penalty_triggered

def calculate_delay_penalties(role):
    """
    Calculates penalties ONLY for Non-Consensual delays.
    Metric: 0.5% deduction per day.
    """
    data = load_complex_data()
    meta = load_full_config().get('meta', {})
    rate = meta.get('cost_settings', {}).get('delay_penalty_rate', 0.5)
    
    delays_log = data.get('delays', [])
    total_deduction_percent = 0.0
    detailed_log = []

    for d in delays_log:
        if d.get('requestor') == role:
            # Check if it was consensual (no penalty) or denied/late (penalty)
            is_consensual = d.get('is_consensual', False)
            status = d.get('status')
            
            # Logic: If denied OR specifically marked non-consensual
            if status == 'Denied' or not is_consensual:
                days = d.get('days', 0)
                if days > 0:
                    penalty = days * rate
                    total_deduction_percent += penalty
                    detailed_log.append(f"{d.get('event', 'Event')} ({days} days late - Non-Consensual)")

    return total_deduction_percent, detailed_log

def analyze_interim_applications(role):
    """Identifies failed applications by this party (Cost Shifting Risk)."""
    data = load_complex_data()
    apps = data.get('applications', [])
    failed_apps = []
    
    for app in apps:
        if app.get('filing_party') == role and app.get('outcome') == 'Denied':
            failed_apps.append(f"{app.get('type')} (Denied on {app.get('date')})")
            
    return failed_apps

def calculate_reversal_amount(offerer_role, offer_date_str):
    """Sums costs incurred by the REJECTOR after the offer date."""
    data = load_complex_data()
    costs = data.get('costs', {})
    
    # If Claimant made offer, Respondent rejected it (and vice versa)
    rejecting_party = 'respondent' if offerer_role == 'claimant' else 'claimant'
    target_log = costs.get(f"{rejecting_party}_log", [])
    
    cutoff_date = datetime.strptime(offer_date_str, "%Y-%m-%d").date()
    reversal_sum = 0.0
    
    for entry in target_log:
        try:
            entry_date = datetime.strptime(entry['date'], "%Y-%m-%d").date()
            if entry_date > cutoff_date:
                reversal_sum += float(entry['amount'])
        except: continue
            
    return reversal_sum, rejecting_party

def get_total_costs(role):
    data = load_complex_data()
    costs = data.get('costs', {}).get(f"{role}_log", [])
    return sum(float(c['amount']) for c in costs)

def check_sealed_offers(final_award_val):
    data = load_complex_data()
    offers = data.get('costs', {}).get('sealed_offers', [])
    reversal_triggers = []
    
    for offer in offers:
        try:
            offer_val = float(offer.get('amount', 0.0))
            # Rule: If Award < Offer, the Offerer is the "effective winner"
            if final_award_val < offer_val:
                reversal_amt, payer = calculate_reversal_amount(offer['offerer'], offer['date'])
                reversal_triggers.append({
                    "offerer": offer['offerer'],
                    "payer": payer,
                    "offer_date": offer['date'],
                    "offer_amount": offer_val,
                    "reversal_sum": reversal_amt
                })
        except: pass
            
    return reversal_triggers

# ==============================================================================
# 2. AI DRAFTING
# ==============================================================================

def try_generate_with_fallback(prompt, project_id, credentials):
    models = ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash-001", "gemini-1.5-pro-001"]
    
    try:
        vertexai.init(project=project_id, location="us-central1", credentials=credentials)
    except Exception as e:
        return f"**[Connection Error]** {e}"

    for m in models:
        try:
            model = GenerativeModel(m)
            response = model.generate_content(prompt)
            return response.text
        except: continue
            
    return "**[System Error]** AI models unavailable."

def generate_cost_award_draft(case_id, final_award_val):
    try:
        # A. GATHER ALL DATA POINTS
        # 1. Totals
        c_total = get_total_costs('claimant')
        r_total = get_total_costs('respondent')
        
        # 2. Conduct (Doc Prod)
        c_score, c_dp_pen = calculate_doc_prod_score('claimant')
        r_score, r_dp_pen = calculate_doc_prod_score('respondent')
        
        # 3. Delays
        c_delay_pct, c_delay_log = calculate_delay_penalties('claimant')
        r_delay_pct, r_delay_log = calculate_delay_penalties('respondent')
        
        # 4. Failed Applications
        c_failed_apps = analyze_interim_applications('claimant')
        r_failed_apps = analyze_interim_applications('respondent')
        
        # 5. Sealed Offers
        reversals = check_sealed_offers(final_award_val) if final_award_val else []
        
        # B. DETAILED NARRATIVE PROMPT
        prompt = f"""
        Act as a TRIBUNAL SECRETARY preparing a "Final Memorandum on Costs" for Case {case_id}.
        
        Your goal is to summarize the procedural history and recommend a cost allocation based STRICTLY on the data below.
        
        --- CASE DATA ---
        
        [A] FINANCIAL BASELINE
        - Claimant Total Request: €{c_total:,.2f}
        - Respondent Total Request: €{r_total:,.2f}
        - Final Principal Award: €{final_award_val:,.2f}
        
        [B] PROCEDURAL HISTORY & CONDUCT
        1. Document Production Phase:
           - Claimant: Rejected {c_score:.1f}% of requests. (Flagged: {c_dp_pen})
           - Respondent: Rejected {r_score:.1f}% of requests. (Flagged: {r_dp_pen})
           *Rule: Rejection >75% constitutes a 'Fishing Expedition' and costs for this phase should be disallowed.*
           
        2. Timeliness & Delays:
           - Claimant Penalties: -{c_delay_pct}% deduction. Events: {c_delay_log}
           - Respondent Penalties: -{r_delay_pct}% deduction. Events: {r_delay_log}
           *Rule: Consensual delays are excused. Non-consensual delays trigger 0.5% deduction/day.*
           
        3. Interim Applications (Pay-as-you-go):
           - Claimant Failed Apps: {c_failed_apps}
           - Respondent Failed Apps: {r_failed_apps}
           *Rule: Loser pays costs of failed interim applications immediately or via offset.*
           
        [C] SETTLEMENT BEHAVIOR (The 'Sealed Offer')
        {f"- CRITICAL: A Sealed Offer was opened. {reversals[0]['offerer']} offered €{reversals[0]['offer_amount']:,.2f} on {reversals[0]['offer_date']}. The Award (€{final_award_val:,.2f}) is LESS favorable. The Rejecting Party ({reversals[0]['payer']}) must pay all costs incurred after the offer date (€{reversals[0]['reversal_sum']:,.2f})." if reversals else "- No cost reversal triggered (Award exceeded all offers or no offers made)."}
        
        --- DRAFTING INSTRUCTIONS ---
        Write a formal legal memorandum (approx 400 words) with the following sections:
        
        I. PROCEDURAL SUMMARY
        Briefly narrate the cost drivers, mentioning the total spend and the disparity between parties.
        
        II. ANALYSIS OF CONDUCT
        Analyze the specific conduct issues above. Explicitly mention if the Claimant's document production ratio was excessive. Discuss the impact of the delays and failed applications.
        
        III. IMPACT OF SEALED OFFERS
        State clearly if the 'Reverse Cost Shifting' rule applies. If so, calculate the impact.
        
        IV. FINAL RECOMMENDATION
        Propose a specific Net Cost Order (e.g., "Respondent shall pay X% of Claimant's costs" or "Claimant shall pay Respondent..."). Deduct penalties and apply reversals mathematically.
        """
        
        if "gcp_service_account" in st.secrets:
            creds = service_account.Credentials.from_service_account_info(st.secrets["gcp_service_account"])
            return try_generate_with_fallback(prompt, st.secrets["gcp_service_account"]["project_id"], creds)
        else:
            return f"**[Demo Mode]** AI Disconnected.\n\nPROMPT:\n{prompt}"

    except Exception as e:
        return f"Error: {e}"

# ==============================================================================
# 3. DOC GENERATOR (UNCHANGED)
# ==============================================================================
def generate_word_document(case_id, draft_text, award_val):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    head = doc.add_heading(f"ARBITRATION CASE: {case_id}", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    title = doc.add_paragraph("FINAL MEMORANDUM ON COSTS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    
    doc.add_paragraph("\n")
    p = doc.add_paragraph("--- PRIVILEGED & CONFIDENTIAL / AI ANALYSIS ---")
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("I. PROCEDURAL CONTEXT", level=1)
    doc.add_paragraph(f"Date: {date.today()}")
    doc.add_paragraph(f"Principal Award: €{award_val:,.2f}")
    
    doc.add_heading("II. RECOMMENDATION", level=1)
    doc.add_paragraph(draft_text)
    
    doc.add_paragraph("\n\n__________________________\nSole Arbitrator")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
